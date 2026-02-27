#!/usr/bin/env python3
"""Daily hidden-discoveries pipeline.

1) Uses OpenAI web search to find 10-15 relevant Earth/Earth-history discovery stories.
2) Creates one citation-backed draft per story in .docx format.
3) Uploads each .docx file to a target Google Drive folder.
"""

from __future__ import annotations

import argparse
import datetime as dt
import json
import os
import re
import socket
import sys
import time
import urllib.error
import urllib.request
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from urllib.parse import urlparse

try:
    from docx import Document
except Exception:
    print("Missing dependency: python-docx. Install with: pip install python-docx", file=sys.stderr)
    raise

try:
    from google.auth.transport.requests import Request
    from google.oauth2 import service_account
    from google.oauth2.credentials import Credentials
    from google_auth_oauthlib.flow import InstalledAppFlow
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaFileUpload
except Exception:
    print(
        "Missing Google API dependencies. Install with: "
        "pip install google-api-python-client google-auth google-auth-oauthlib",
        file=sys.stderr,
    )
    raise

OPENAI_API_URL = "https://api.openai.com/v1/responses"
DRIVE_SCOPES = ["https://www.googleapis.com/auth/drive"]
BLOCKED_DISCOVERY_DOMAINS = {
    "chasewaterdogs.co.uk",
}


@dataclass
class Story:
    title: str
    summary: str
    source: str
    date: str
    url: str
    category: str


@dataclass
class Draft:
    title: str
    intro: str
    sections: list[dict[str, Any]]
    closing: str
    closing_citation_url: str


@dataclass
class StoryResearch:
    story_title: str
    evidence: list[dict[str, str]]


def load_env_file(path: Path) -> None:
    if not path.exists():
        return
    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        if key and key not in os.environ:
            os.environ[key] = value


def require_env(name: str) -> str:
    value = os.getenv(name, "").strip()
    if not value:
        raise RuntimeError(f"Missing required env var: {name}")
    return value


def env_bool(name: str, default: bool) -> bool:
    value = os.getenv(name)
    if value is None:
        return default
    return value.strip().lower() in {"1", "true", "yes", "on"}


def env_int(name: str, default: int) -> int:
    value = os.getenv(name)
    if value is None or not value.strip():
        return default
    return int(value)


def response_output_text(payload: dict[str, Any]) -> str:
    if isinstance(payload.get("output_text"), str):
        return payload["output_text"]

    chunks: list[str] = []
    for item in payload.get("output", []):
        for content in item.get("content", []):
            if content.get("type") in {"output_text", "text"} and content.get("text"):
                chunks.append(content["text"])
    return "\n".join(chunks).strip()


def call_openai(
    api_key: str,
    model: str,
    prompt: str,
    schema_name: str,
    schema: dict[str, Any],
    use_web_search: bool,
) -> dict[str, Any]:
    body: dict[str, Any] = {
        "model": model,
        "input": [{"role": "user", "content": [{"type": "input_text", "text": prompt}]}],
        "text": {
            "format": {
                "type": "json_schema",
                "name": schema_name,
                "strict": True,
                "schema": schema,
            }
        },
    }
    if use_web_search:
        body["tools"] = [{"type": "web_search_preview"}]

    req = urllib.request.Request(
        OPENAI_API_URL,
        data=json.dumps(body).encode("utf-8"),
        method="POST",
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
    )

    timeout_seconds = env_int("OPENAI_TIMEOUT_SECONDS", 300)
    max_retries = env_int("OPENAI_MAX_RETRIES", 3)
    last_exc: Exception | None = None
    payload: dict[str, Any] | None = None
    for attempt in range(1, max_retries + 1):
        try:
            with urllib.request.urlopen(req, timeout=timeout_seconds) as resp:
                payload = json.loads(resp.read().decode("utf-8"))
            break
        except urllib.error.HTTPError as err:
            detail = err.read().decode("utf-8", errors="replace")
            raise RuntimeError(f"OpenAI API error ({err.code}): {detail}") from err
        except (urllib.error.URLError, socket.timeout, TimeoutError) as err:
            last_exc = err
            if attempt == max_retries:
                break
            time.sleep(min(20, 2 * attempt))

    if payload is None:
        raise RuntimeError(f"OpenAI API request failed after {max_retries} attempts: {last_exc}")

    text = response_output_text(payload)
    if not text:
        raise RuntimeError("OpenAI response did not contain text output.")

    try:
        return json.loads(text)
    except json.JSONDecodeError as exc:
        raise RuntimeError(f"Model did not return valid JSON. Output: {text[:500]}") from exc


def build_story_prompt(
    min_count: int, max_count: int, lookback_days: int, excluded_urls: list[str] | None = None
) -> str:
    today = dt.date.today()
    start_date = today - dt.timedelta(days=lookback_days)
    window_note = (
        "Treat this as a 3-month lookback window."
        if lookback_days == 90
        else f"Treat this as a {lookback_days}-day lookback window."
    )
    exclusion_note = ""
    if excluded_urls:
        short = excluded_urls[:20]
        exclusion_note = (
            "\nDo not include these URLs again:\n- " + "\n- ".join(short)
        )
    return f"""
Give me a comprehensive list of discovery stories about something discovered hidden under, within, buried beneath, or submerged under something else.

Include stories where:
- things were found under the ocean, under ice sheets, beneath the surface of Earth, within mountains, buried deep in sediments, inside caves or underground chambers, or submerged archaeological sites;
and stories where discoveries reveal hidden ancient ecosystems, buried landscapes, fossils under ice or rock, underwater cities/villages, formerly submerged landforms, unknown subterranean features, sealed archaeological chambers, deep Earth hidden reservoirs, fossils that rewrite history, or hidden climate/biological records.

For each story, provide: title/headline, 1-2 sentence summary, source publication, date, direct URL, and category.
Aim for {min_count}-{max_count} distinct discoveries across archaeology, geology, paleontology, and environmental science.
Hard constraints: Earth/Earth-history only, no generic astronomy, and strong fit to hidden-beneath/within/under theme.
Time constraint: include ONLY stories published between {start_date.isoformat()} and {today.isoformat()} (last {lookback_days} days). Exclude anything older.
{window_note}
Source quality constraint:
- Prefer reputable science/news/academic/government sources.
- Avoid low-credibility, content-farm, or spam domains.
{exclusion_note}
""".strip()


def discover_stories(
    api_key: str, model: str, min_count: int, max_count: int, lookback_days: int
) -> list[Story]:
    schema = {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "stories": {
                "type": "array",
                "minItems": min_count,
                "maxItems": max_count,
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "title": {"type": "string"},
                        "summary": {"type": "string"},
                        "source": {"type": "string"},
                        "date": {"type": "string"},
                        "url": {"type": "string"},
                        "category": {
                            "type": "string",
                            "enum": [
                                "archaeology",
                                "geology",
                                "paleontology",
                                "environmental_science",
                                "climate_record",
                                "mixed",
                            ],
                        },
                    },
                    "required": ["title", "summary", "source", "date", "url", "category"],
                },
            }
        },
        "required": ["stories"],
    }

    max_attempts = env_int("DISCOVERY_MAX_RETRIES", 5)
    min_domains = env_int("DISCOVERY_MIN_DISTINCT_DOMAINS", 4)
    by_url: dict[str, Story] = {}

    for _ in range(max_attempts):
        excluded_urls = list(by_url.keys())
        payload = call_openai(
            api_key=api_key,
            model=model,
            prompt=build_story_prompt(min_count, max_count, lookback_days, excluded_urls),
            schema_name="hidden_discoveries",
            schema=schema,
            use_web_search=True,
        )
        candidate = [
            Story(
                title=row["title"].strip(),
                summary=row["summary"].strip(),
                source=row["source"].strip(),
                date=row["date"].strip(),
                url=row["url"].strip(),
                category=row["category"].strip(),
            )
            for row in payload["stories"]
        ]

        filtered: list[Story] = []
        today = dt.date.today()
        min_date = today - dt.timedelta(days=lookback_days)
        for item in candidate:
            domain = (urlparse(item.url).netloc or "").lower()
            if domain in BLOCKED_DISCOVERY_DOMAINS:
                continue
            if not item.url.startswith("http"):
                continue
            try:
                published = dt.date.fromisoformat(item.date.strip())
            except Exception:
                continue
            if not (min_date <= published <= today):
                continue
            filtered.append(item)

        distinct_domains = len({(urlparse(s.url).netloc or "").lower() for s in filtered})
        for story in filtered:
            by_url[story.url] = story

        aggregate = list(by_url.values())[:max_count]
        agg_domains = len({(urlparse(s.url).netloc or "").lower() for s in aggregate})
        if len(aggregate) >= min_count and agg_domains >= min_domains:
            return aggregate

    aggregate = list(by_url.values())[:max_count]
    require_min = env_bool("DISCOVERY_REQUIRE_MIN_COUNT", False)
    if require_min and len(aggregate) < min_count:
        raise RuntimeError(
            f"Discovery underfilled after retries: got {len(aggregate)} stories, need at least {min_count}."
        )
    return aggregate


def build_research_prompt(story: Story, lookback_days: int) -> str:
    today = dt.date.today()
    start_date = today - dt.timedelta(days=lookback_days)
    return f"""
Research this discovery story with web search and compile a structured evidence pack.

Story seed:
- Title: {story.title}
- Summary: {story.summary}
- Source: {story.source}
- Date: {story.date}
- URL: {story.url}
- Category: {story.category}

Requirements:
- Find corroborating and contextual sources, not just one article.
- Focus on sources published between {start_date.isoformat()} and {today.isoformat()} where possible.
- Build 8 to 10 evidence entries.
- Each evidence entry must have:
  - angle (what this point covers),
  - paragraph (55-95 words),
  - citation_url (direct link),
  - source (publisher/site),
  - date (YYYY-MM-DD if available).
- Use at least 6 distinct citation URLs and 4 distinct source domains.
- Avoid repeating the same URL many times.
- Keep strictly Earth/Earth-history context.
""".strip()


def research_story(api_key: str, model: str, story: Story, lookback_days: int) -> StoryResearch:
    schema = {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "story_title": {"type": "string"},
            "evidence": {
                "type": "array",
                "minItems": 8,
                "maxItems": 10,
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "angle": {"type": "string"},
                        "paragraph": {"type": "string"},
                        "citation_url": {"type": "string"},
                        "source": {"type": "string"},
                        "date": {"type": "string"},
                    },
                    "required": ["angle", "paragraph", "citation_url", "source", "date"],
                },
            },
        },
        "required": ["story_title", "evidence"],
    }
    max_attempts = env_int("RESEARCH_MAX_RETRIES", 3)
    last_exc: Exception | None = None
    payload: dict[str, Any] | None = None
    for _ in range(max_attempts):
        try:
            payload = call_openai(
                api_key=api_key,
                model=model,
                prompt=build_research_prompt(story, lookback_days),
                schema_name="story_research_pack",
                schema=schema,
                use_web_search=True,
            )
            break
        except RuntimeError as exc:
            last_exc = exc
            time.sleep(2)
    if payload is None:
        raise RuntimeError(f"Research generation failed after {max_attempts} attempts: {last_exc}")
    return StoryResearch(
        story_title=payload["story_title"].strip(),
        evidence=payload["evidence"],
    )


def build_draft_prompt(story: Story, research: StoryResearch) -> str:
    evidence_json = json.dumps(research.evidence, ensure_ascii=True)
    return f"""
Create a professional article draft for this story:
Title: {story.title}
Summary: {story.summary}
Source: {story.source}
Date: {story.date}
URL: {story.url}
Category: {story.category}

Output structure:
1) Title
2) Intro paragraph (140-220 words)
3) 6 to 8 sections. Each section must include:
   - heading
   - 2 paragraphs (each 110-170 words)
   - one citation URL for each paragraph
4) Closing paragraph (120-180 words) + citation URL

Rules:
- Use ONLY the research evidence provided below.
- Each paragraph must be factual and source-grounded.
- Put citation URL directly below each paragraph (as a URL string).
- Keep an analytical and nuanced tone, not generic.
- Stay on Earth/Earth-history context only.
- Use multiple independent sources, not a single source rewrite.
- Citation quality constraints:
  - Every paragraph must have its own citation URL.
  - Avoid reusing the same citation URL for multiple paragraphs.
  - Use at least 5 distinct citation URLs across the draft.
  - Use at least 3 distinct source domains across the draft.
  - Do not use the original story URL for every paragraph.

Research evidence JSON:
{evidence_json}
""".strip()


def generate_draft(api_key: str, model: str, story: Story, research: StoryResearch) -> Draft:
    schema = {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "title": {"type": "string"},
            "intro": {"type": "string"},
            "sections": {
                "type": "array",
                "minItems": 6,
                "maxItems": 8,
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "heading": {"type": "string"},
                        "paragraphs": {
                            "type": "array",
                            "minItems": 2,
                            "maxItems": 2,
                            "items": {
                                "type": "object",
                                "additionalProperties": False,
                                "properties": {
                                    "text": {"type": "string"},
                                    "citation_url": {"type": "string"},
                                },
                                "required": ["text", "citation_url"],
                            },
                        },
                    },
                    "required": ["heading", "paragraphs"],
                },
            },
            "closing": {"type": "string"},
            "closing_citation_url": {"type": "string"},
        },
        "required": ["title", "intro", "sections", "closing", "closing_citation_url"],
    }

    payload = call_openai(
        api_key=api_key,
        model=model,
        prompt=build_draft_prompt(story, research),
        schema_name="story_draft",
        schema=schema,
        use_web_search=False,
    )

    return Draft(
        title=payload["title"].strip(),
        intro=payload["intro"].strip(),
        sections=payload["sections"],
        closing=payload["closing"].strip(),
        closing_citation_url=payload["closing_citation_url"].strip(),
    )


def collect_citation_urls(draft: Draft) -> list[str]:
    urls: list[str] = []
    for section in draft.sections:
        for para in section.get("paragraphs", []):
            url = str(para.get("citation_url", "")).strip()
            if url:
                urls.append(url)
    closing = draft.closing_citation_url.strip()
    if closing:
        urls.append(closing)
    return urls


def citation_domain(url: str) -> str:
    try:
        return (urlparse(url).netloc or "").lower()
    except Exception:
        return ""


def draft_meets_citation_quality(draft: Draft, story_url: str) -> bool:
    urls = collect_citation_urls(draft)
    if not urls:
        return False

    unique_urls = set(urls)
    unique_domains = {citation_domain(u) for u in unique_urls if citation_domain(u)}
    story_url_count = sum(1 for u in urls if u.strip() == story_url.strip())

    # Require high diversity to avoid single-source rewrites.
    if len(unique_urls) < 5:
        return False
    if len(unique_domains) < 3:
        return False
    if story_url_count >= len(urls):
        return False
    return True


def generate_draft_with_retries(api_key: str, model: str, story: Story, max_attempts: int) -> Draft:
    raise RuntimeError("Deprecated: use generate_draft_with_research_retries")


def generate_draft_with_research_retries(
    api_key: str,
    model: str,
    story: Story,
    research: StoryResearch,
    max_attempts: int,
) -> Draft:
    last_draft: Draft | None = None
    allowed_urls = {
        str(item.get("citation_url", "")).strip()
        for item in research.evidence
        if str(item.get("citation_url", "")).strip()
    }
    for _ in range(max_attempts):
        candidate = generate_draft(api_key, model, story, research)
        last_draft = candidate
        candidate_urls = collect_citation_urls(candidate)
        urls_from_research = all((u in allowed_urls) for u in candidate_urls if u)
        if draft_meets_citation_quality(candidate, story.url) and urls_from_research:
            return candidate
    if last_draft is None:
        raise RuntimeError("Draft generation failed with no output.")
    return last_draft


def sanitize_filename(name: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9 ._-]+", "", name).strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    return (cleaned or "story")[:120]


def write_story_docx(output_dir: Path, story: Story, draft: Draft) -> Path:
    timestamp = dt.datetime.now().strftime("%Y-%m-%d")
    filename = f"{timestamp} - {sanitize_filename(story.title)}.docx"
    path = output_dir / filename

    document = Document()
    document.add_heading(draft.title, level=1)

    document.add_paragraph(draft.intro)
    document.add_paragraph(story.url)

    for section in draft.sections:
        document.add_heading(section["heading"].strip(), level=2)
        for para in section["paragraphs"]:
            text = para.get("text", "").strip()
            citation = para.get("citation_url", "").strip() or story.url
            if text:
                document.add_paragraph(text)
            document.add_paragraph(citation)

    document.add_paragraph(draft.closing)
    document.add_paragraph(draft.closing_citation_url or story.url)

    document.save(path)
    return path


def write_prompt_response_docx(output_dir: Path, stories: list[Story]) -> Path:
    date_stamp = dt.datetime.now().strftime("%Y-%m-%d")
    path = output_dir / f"{date_stamp}_prompt_respoonse.docx"
    document = Document()
    document.add_heading("Prompt Response - Discovery Stories", level=1)
    document.add_paragraph(f"Generated on: {dt.datetime.now().isoformat(timespec='seconds')}")
    document.add_paragraph(f"Total stories: {len(stories)}")

    for idx, story in enumerate(stories, start=1):
        document.add_heading(f"{idx}. {story.title}", level=2)
        document.add_paragraph(story.summary)
        document.add_paragraph(f"Source: {story.source}")
        document.add_paragraph(f"Date: {story.date}")
        document.add_paragraph(f"Category: {story.category}")
        document.add_paragraph(story.url)

    document.save(path)
    return path


def get_drive_service(repo_root: Path):
    auth_mode = os.getenv("GOOGLE_DRIVE_AUTH_MODE", "service_account").strip().lower()
    if auth_mode == "oauth":
        client_secret_file = require_env("GOOGLE_OAUTH_CLIENT_SECRET_FILE")
        token_file = Path(
            os.getenv(
                "GOOGLE_OAUTH_TOKEN_FILE",
                str(repo_root / "secrets" / "google_oauth_token.json"),
            )
        )

        creds: Credentials | None = None
        if token_file.exists():
            creds = Credentials.from_authorized_user_file(str(token_file), DRIVE_SCOPES)

        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())

        if not creds or not creds.valid:
            flow = InstalledAppFlow.from_client_secrets_file(client_secret_file, DRIVE_SCOPES)
            creds = flow.run_local_server(port=0, access_type="offline", prompt="consent")
            token_file.parent.mkdir(parents=True, exist_ok=True)
            token_file.write_text(creds.to_json(), encoding="utf-8")
            print(f"OAuth token saved: {token_file}")

        return build("drive", "v3", credentials=creds)

    service_account_file = require_env("GOOGLE_SERVICE_ACCOUNT_FILE")
    creds = service_account.Credentials.from_service_account_file(
        service_account_file, scopes=DRIVE_SCOPES
    )
    return build("drive", "v3", credentials=creds)


def maybe_delete_existing(service: Any, folder_id: str, filename: str) -> None:
    safe_name = filename.replace("'", "\\'")
    query = f"name = '{safe_name}' and '{folder_id}' in parents and trashed = false"
    resp = (
        service.files()
        .list(q=query, fields="files(id,name)", supportsAllDrives=True, includeItemsFromAllDrives=True)
        .execute()
    )
    for item in resp.get("files", []):
        service.files().delete(fileId=item["id"], supportsAllDrives=True).execute()


def upload_to_drive(service: Any, folder_id: str, file_path: Path, overwrite_existing: bool) -> str:
    if overwrite_existing:
        maybe_delete_existing(service, folder_id, file_path.name)

    media = MediaFileUpload(
        str(file_path),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        resumable=False,
    )
    metadata = {"name": file_path.name, "parents": [folder_id]}
    created = (
        service.files()
        .create(body=metadata, media_body=media, fields="id,name,webViewLink", supportsAllDrives=True)
        .execute()
    )
    return created.get("webViewLink", "")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Hidden discoveries daily automation")
    parser.add_argument(
        "--auth-only",
        action="store_true",
        help="Run only Google Drive auth setup and exit.",
    )
    return parser.parse_args()


def main(auth_only: bool = False) -> int:
    repo_root = Path(__file__).resolve().parents[1]
    load_env_file(repo_root / ".env.hidden_discoveries")

    api_key = require_env("OPENAI_API_KEY")
    model = os.getenv("OPENAI_MODEL", "gpt-4.1-mini").strip()
    discovery_model = os.getenv("DISCOVERY_MODEL", model).strip()
    research_model = os.getenv("RESEARCH_MODEL", discovery_model).strip()
    draft_model = os.getenv("DRAFT_MODEL", model).strip()
    folder_id = require_env("GOOGLE_DRIVE_FOLDER_ID")

    min_count = env_int("STORY_COUNT_MIN", 10)
    max_count = env_int("STORY_COUNT_MAX", 15)
    lookback_days = env_int("LOOKBACK_DAYS", 7)
    overwrite_existing = env_bool("OVERWRITE_EXISTING", False)
    draft_max_citation_retries = env_int("DRAFT_MAX_CITATION_RETRIES", 3)
    output_dir = Path(os.getenv("OUTPUT_DIR", str(repo_root / "outputs" / "hidden_discoveries")))
    output_dir.mkdir(parents=True, exist_ok=True)

    print("[0/4] Connecting to Google Drive...")
    drive_service = get_drive_service(repo_root)
    if auth_only:
        # Basic folder read validates auth and folder access.
        drive_service.files().get(
            fileId=folder_id,
            fields="id,name",
            supportsAllDrives=True,
        ).execute()
        print("Google Drive auth check passed.")
        return 0

    print(
        f"[1/5] Discovering stories ({min_count}-{max_count}) via OpenAI web search "
        f"(last {lookback_days} days)..."
    )
    stories = discover_stories(api_key, discovery_model, min_count, max_count, lookback_days)
    print(f"Found {len(stories)} stories.")
    prompt_response_file = write_prompt_response_docx(output_dir, stories)

    print("[2/5] Researching each story with a second web search pass...")
    research_by_index: dict[int, StoryResearch] = {}
    for idx, story in enumerate(stories, start=1):
        print(f"  - ({idx}/{len(stories)}) Researching: {story.title}")
        research_by_index[idx] = research_story(
            api_key=api_key,
            model=research_model,
            story=story,
            lookback_days=lookback_days,
        )

    print("[3/5] Generating drafts from researched evidence...")
    generated_files: list[Path] = []
    for idx, story in enumerate(stories, start=1):
        print(f"  - ({idx}/{len(stories)}) {story.title}")
        draft = generate_draft_with_research_retries(
            api_key=api_key,
            model=draft_model,
            story=story,
            research=research_by_index[idx],
            max_attempts=draft_max_citation_retries,
        )
        file_path = write_story_docx(output_dir, story, draft)
        generated_files.append(file_path)
    generated_files.append(prompt_response_file)

    print("[4/5] Uploading files...")
    for path in generated_files:
        link = upload_to_drive(
            service=drive_service,
            folder_id=folder_id,
            file_path=path,
            overwrite_existing=overwrite_existing,
        )
        if link:
            print(f"  - Uploaded: {path.name} -> {link}")
        else:
            print(f"  - Uploaded: {path.name}")

    print("Done.")
    return 0


if __name__ == "__main__":
    args = parse_args()
    raise SystemExit(main(auth_only=args.auth_only))
